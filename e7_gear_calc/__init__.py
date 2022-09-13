from PIL import Image
import pyautogui
import pytesseract
import cv2
import numpy as np
from docx import Document
from docx.shared import Inches
import io

def record_output(output_text, img):
    width, height = img.size
    img = img.resize((int(width/4), int(height/4)))

    try:
        document = Document("output.docx")
    except:
        document = Document()

    p = document.add_paragraph()
    r = p.add_run()
    image_file = io.BytesIO()
    img.save(image_file, format="PNG")
    r.add_picture(image_file)
    r.add_text(output_text)

    document.save("output.docx")

def get_rows(str):
    rows = str.split('\n')
    rows.pop()
    
    output = []
    index = 0

    for row in rows:
        if len(row) > 3:
            output.append(row)
        elif len(row) >= 1:
            output[index] = output[index] + row
            index += 1            

    return output

def get_info(row):
    return [''.join([i for i in row if not i.isdigit() and i != " "]), ''.join([i for i in row if i.isdigit()])]

def filter_image(image):
    hsv = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)

    lower = np.array([0, 0, 83])
    upper = np.array([0, 0, 200])

    filteredImg = cv2.inRange(hsv, lower, upper)

    return filteredImg

def get_gear_score_from_image(image, filter=True):
    multipliers = {
        "Attack%": 1,
        "Health%": 1,
        "Defense%": 1,
        "EffectResistance%": 1,
        "Effectiveness%": 1,
        "Attack": (1/12),
        "Health": (1/60),
        "Defense": (1/6),
        "CriticalHitChance%": 1.6,
        "criticalHitChance%": 1.6,
        "CriticalHitDamage%": 1.1,
        "criticalHitDamage%": 1.1,
        "Speed": 1.8
    }

    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract'
    rows = []

    try:
        gearImage = np.array(image)

        if filter:
            gearImage = filter_image(gearImage)

        result = pytesseract.image_to_string(gearImage, config="-c tessedit_char_whitelist=%0123456789pEHgeRsvdnSaickrClmhfAtD")   
        
        rows = get_rows(result)
        output = "\n"
        score = 0

        for row in rows:
            stat_name, statValue = get_info(row)
            if stat_name not in multipliers:
                if stat_name[-1:] == "%":
                    stat_name = stat_name[:-2] + stat_name[-1:]
                else:
                    stat_name = stat_name[:-1]
            score += multipliers[stat_name] * int(statValue)
            output += (stat_name + ": " + str(statValue) + "\n")
        
        output += ("Score: " + str(score) + "\n")
        
        return output

    except Exception as e:
        print("An error occurred: " + str(e))
        print(rows)
        print(result)
        cv2.imshow('image', gearImage)
        cv2.waitKey(0)
        return "Error"
    
def main():
    mainImage = pyautogui.screenshot(region=(680, 223, 720, 770))
    statsImage = pyautogui.screenshot(region=(680, 773, 650, 250))

    output = str(get_gear_score_from_image(statsImage))

    if output != "Error":
        print(output)

        # Record the result into a document.
        record_output(output, mainImage)

main()
